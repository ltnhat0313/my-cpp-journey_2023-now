from docx import Document

def create_report(filename):
    doc = Document()
    
    # Tiêu đề
    doc.add_heading('Báo cáo ARP và CAM Table Overflow', level=1)
    
    # Phần 1: Giao thức ARP và lỗ hổng ARP Spoofing
    doc.add_heading('1. Giao thức ARP và ARP Spoofing', level=2)
    doc.add_paragraph(
        "ARP (Address Resolution Protocol) là giao thức chuyển đổi địa chỉ IP thành địa chỉ MAC. "
        "Do không có cơ chế xác thực, ARP dễ bị tấn công ARP Spoofing, trong đó kẻ tấn công gửi "
        "gói ARP giả mạo để ánh xạ địa chỉ IP của một thiết bị hợp lệ với địa chỉ MAC của kẻ tấn công. "
        "Kết quả là lưu lượng mạng của thiết bị hợp lệ sẽ bị chuyển hướng đến kẻ tấn công."
    )
    
    # Phần 2: Cơ chế học của Switch và bảng CAM
    doc.add_heading('2. Cơ chế học của Switch và bảng CAM', level=2)
    doc.add_paragraph(
        "Switch học địa chỉ MAC của thiết bị bằng cách ghi lại cổng nhận được khung dữ liệu. "
        "Bảng CAM lưu ánh xạ giữa địa chỉ MAC và cổng. Nếu bảng CAM bị tràn do tấn công (CAM Table Overflow), "
        "switch sẽ hoạt động như một hub và phát dữ liệu ra tất cả các cổng, tạo điều kiện cho nghe lén và tấn công."
    )
    
    # Phần 3: Nguy cơ và giải pháp phòng ngừa
    doc.add_heading('3. Nguy cơ và giải pháp phòng ngừa', level=2)
    doc.add_paragraph(
        "Nguy cơ:\n- ARP Spoofing: Dẫn đến nghe lén và tấn công MITM.\n"
        "- CAM Table Overflow: Cho phép kẻ tấn công nghe lén lưu lượng mạng.\n\n"
        "Giải pháp phòng ngừa:\n- Sử dụng ARP tĩnh và công cụ phát hiện ARP Spoofing.\n"
        "- Cấu hình Port Security trên Switch để giới hạn số lượng địa chỉ MAC mỗi cổng.\n"
        "- Phân chia mạng thành VLAN để hạn chế phạm vi tấn công."
    )
    
    doc.save(filename)
    print(f"File Word '{filename}' đã được tạo thành công!")

# Gọi hàm tạo báo cáo
create_report("BaoCao_ARP_CAM.docx")
